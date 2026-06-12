import os
import sys
import cProfile
import pstats
import io
import platform
from datetime import datetime
from docx import Document

os.environ['QT_QPA_PLATFORM'] = 'xcb'

report_name = 'profile_report.docx'
profile_file = 'profile_startup.prof'

profiler = cProfile.Profile()
profiler.enable()

# Import modules and run database initialization as a startup simulation.
from db import initialize_database
from utils import resource_path, text_to_speech
from core import check_product

initialize_database('DB_FILE')

profiler.disable()
profiler.dump_stats(profile_file)

text_stream = io.StringIO()
stats = pstats.Stats(profiler, stream=text_stream).sort_stats('cumulative')
stats.print_stats(30)
raw_stats = text_stream.getvalue()

# Build the report document.
doc = Document()
doc.add_heading('Profile Report', level=1)
doc.add_paragraph(f'Date: {datetime.now().strftime("%d-%m-%Y %I:%M:%S %p")}')
doc.add_paragraph(f'Python: {platform.python_version()} ({platform.python_implementation()})')
doc.add_paragraph(f'Platform: {platform.system()} {platform.release()}')
doc.add_paragraph('Profile target: module import and database initialization startup flow.')

doc.add_heading('Top 30 Functions by Cumulative Time', level=2)
for line in raw_stats.splitlines():
    doc.add_paragraph(line)

doc.add_heading('Notes', level=2)
doc.add_paragraph('This report profiles the import path and database initialization of the inventory management application. It does not execute the GUI main loop.')
doc.add_paragraph('The most expensive startup operations are likely module imports and database schema setup.')

doc.save(report_name)
print(f'Report generated: {report_name}')
